"""
extract_diagrams_docx.py — DOCX diagram extraction v5

CHANGES in v5 (FINAL):
  - Chemistry: DUMP MODE — extract ALL images per question into
    question_diagram_urls bag. No zone detection, no junk filter.
    Reviewer assigns zones manually via dashboard dropdown.
  - Biology: ZONE MODE (unchanged from v4) — element-order zone
    detection with redistribution logic.
  - Junk filter DISABLED for Chemistry (reviewer labels junk in dashboard).
  - Junk filter ENABLED for Biology (auto-filter still useful).

WHY:
  iLovePDF-converted DOCX has unreliable element ordering for Chemistry
  questions. Zone detection assigns images to wrong options. Missing
  images because junk filter rejects legitimate chemical structures.
  Solution: extract everything, let human review.
"""

import re
import io
import zipfile
import tempfile
import os
from collections import defaultdict

from docx import Document
from docx.oxml.ns import qn as docx_qn
from PIL import Image
import numpy as np
from google.cloud import storage

PROJECT         = "project-3639c8e1-b432-4a18-99f"
DIAGRAMS_BUCKET = f"{PROJECT}-diagrams"

MIN_IMAGE_BYTES      = 200
WATERMARK_THRESHOLD  = 10
JUNK_WHITE_THRESHOLD = 245
JUNK_MIN_DARK_PIXELS = 100
JUNK_MIN_DARK_RATIO  = 0.005

# Chemistry uses DUMP mode (all images → bag, no zone detection)
# Biology uses ZONE mode (element-order zone detection)
DUMP_SECTIONS = {"Chemistry"}
ZONE_SECTIONS = {"Biology"}
DOCX_SECTIONS = DUMP_SECTIONS | ZONE_SECTIONS

METADATA_PREFIXES = [
    'Question Type', 'Difficulty', 'Topic', 'Concept',
    'Single Correct', 'Expected', 'NEET UG',
]


def upload_to_gcs(image_bytes, filename, paper_id):
    sc = storage.Client(project=PROJECT)
    bp = f"{paper_id}/{filename}"
    sc.bucket(DIAGRAMS_BUCKET).blob(bp).upload_from_string(
        image_bytes, content_type="image/png")
    return f"gs://{DIAGRAMS_BUCKET}/{bp}"


def ensure_png(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode == 'RGBA':
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        out = io.BytesIO()
        img.save(out, format='PNG', optimize=True)
        return out.getvalue()
    except Exception:
        return image_bytes


def is_junk_image(image_bytes):
    """Check actual pixel content — filters empty rectangles and dots."""
    try:
        img = Image.open(io.BytesIO(image_bytes)).convert('RGB')
        arr = np.array(img)
        total = arr.shape[0] * arr.shape[1]
        if arr.shape[0] < 10 or arr.shape[1] < 10:
            return True
        dark = int(np.sum(arr.mean(axis=2) < JUNK_WHITE_THRESHOLD))
        if dark < JUNK_MIN_DARK_PIXELS:
            return True
        if dark / total < JUNK_MIN_DARK_RATIO:
            return True
        if total > 400 and arr.shape[0] > 6 and arr.shape[1] > 6:
            interior = arr[3:-3, 3:-3]
            int_dark = int(np.sum(interior.mean(axis=2) < JUNK_WHITE_THRESHOLD))
            if interior.shape[0] * interior.shape[1] > 0:
                if int_dark / (interior.shape[0] * interior.shape[1]) < 0.001:
                    return True
        return False
    except Exception:
        return False


def detect_watermark_rids(doc):
    rid_count = defaultdict(int)
    for para in doc.paragraphs:
        seen = set()
        for blip in para._element.findall('.//' + docx_qn('a:blip')):
            rid = blip.get(docx_qn('r:embed'))
            if rid and rid not in seen:
                rid_count[rid] += 1
                seen.add(rid)
    wm = {rid for rid, c in rid_count.items() if c >= WATERMARK_THRESHOLD}
    if wm:
        print(f"  Watermark rIds: {wm}")
    return wm


def build_rid_to_image(docx_path):
    rid_to_image = {}
    with zipfile.ZipFile(docx_path) as z:
        rels = z.read("word/_rels/document.xml.rels").decode()
        for m in re.finditer(r'Id="(rId\d+)".*?Target="(media/[^"]+)"', rels):
            rid, media = m.group(1), m.group(2)
            try:
                data = z.read(f"word/{media}")
                if len(data) >= MIN_IMAGE_BYTES:
                    rid_to_image[rid] = {"bytes": data, "file": media}
            except Exception:
                pass
    print(f"  Valid images: {len(rid_to_image)}")
    return rid_to_image


def detect_section_change(text):
    if text == "Physics":
        return "Physics"
    if "NEET" in text and "Chemistry" in text and len(text) < 80:
        return "Chemistry"
    if re.match(r'^Biology\s*\d*$', text):
        return "Biology"
    return None


def is_metadata(text):
    return any(text.startswith(p) for p in METADATA_PREFIXES)


def para_has_sol(para):
    for child in para._element.iter():
        if child.tag.endswith('}t'):
            t = (child.text or '').strip()
            if t.startswith('Sol'):
                return True
    return False


def extract_diagrams_docx(docx_bytes, paper_id, corrected_questions=None):
    """
    v5 extraction:
      Chemistry: DUMP mode — all images per question → question bag
      Biology: ZONE mode — element-order zone detection (v4 logic)
    """
    print(f"\n{'='*60}")
    print(f"DOCX EXTRACTION v5 (dump+zone) — {paper_id}")
    print(f"{'='*60}")

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.write(docx_bytes)
    tmp.close()
    docx_path = tmp.name

    try:
        doc = Document(docx_path)
        print(f"  Paragraphs: {len(doc.paragraphs)}")

        watermark_rids = detect_watermark_rids(doc)
        rid_to_image = build_rid_to_image(docx_path)

        # Build question lists per section from Gemini data
        section_questions = {"Physics": [], "Chemistry": [], "Biology": []}
        if corrected_questions:
            for q in corrected_questions:
                sec = q.get("section", "")
                if sec in section_questions:
                    section_questions[sec].append(q)
            for sec in section_questions:
                section_questions[sec].sort(
                    key=lambda q: q.get("question_number", 0))

        # ─── PASS 1: Find section boundaries and Sol. positions ───

        sol_positions = []
        section_ranges = {}
        current_section = "Physics"
        section_start = 0

        for pi, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            new_sec = detect_section_change(text)
            if new_sec and new_sec != current_section:
                section_ranges[current_section] = (section_start, pi)
                current_section = new_sec
                section_start = pi
            if para_has_sol(para):
                sol_positions.append((pi, current_section))

        section_ranges[current_section] = (section_start, len(doc.paragraphs))

        for sec, (s, e) in section_ranges.items():
            tag = "→ DUMP" if sec in DUMP_SECTIONS else "→ ZONE" if sec in ZONE_SECTIONS else "→ SKIP"
            sec_sols = sum(1 for _, sc in sol_positions if sc == sec)
            qs = section_questions.get(sec, [])
            if qs:
                nums = [q["question_number"] for q in qs]
                print(f"  {sec}: P{s}-P{e}, {sec_sols} Sol., "
                      f"Q{min(nums)}-Q{max(nums)} {tag}")

        # ─── PASS 2: Process each section ───

        url_map = defaultdict(lambda: defaultdict(list))
        images_uploaded = 0
        junk_filtered = 0

        for section in DOCX_SECTIONS:
            if section not in section_ranges:
                continue

            sec_start, sec_end = section_ranges[section]
            sec_sols = [pi for pi, sc in sol_positions if sc == section]
            sec_qs = section_questions.get(section, [])

            if not sec_qs:
                continue

            is_dump = section in DUMP_SECTIONS
            mode_label = "DUMP" if is_dump else "ZONE"
            print(f"\n  Processing {section} ({mode_label}): "
                  f"{len(sec_sols)} Sol., {len(sec_qs)} questions")

            for qi in range(min(len(sec_sols), len(sec_qs))):
                q_start = sec_sols[qi - 1] + 1 if qi > 0 else sec_start
                q_end = sec_sols[qi]  # Sol. paragraph (inclusive)
                q_num = sec_qs[qi]["question_number"]
                q_sec = sec_qs[qi]["section"]
                key = f"{q_sec}_{q_num}"

                if is_dump:
                    # ══════════════════════════════════════════
                    # DUMP MODE (Chemistry)
                    # Extract ALL images → question bag
                    # No zone detection, no junk filter
                    # Solution images still separated
                    # ══════════════════════════════════════════

                    for pi in range(q_start, q_end + 1):
                        if pi >= len(doc.paragraphs):
                            break
                        para = doc.paragraphs[pi]
                        is_sol_para = para_has_sol(para)

                        for child in para._element.iter():
                            if child.tag.endswith('}blip'):
                                rid = child.get(docx_qn('r:embed'))
                                if (rid and rid not in watermark_rids
                                        and rid in rid_to_image):
                                    img_data = rid_to_image[rid]
                                    img_bytes = ensure_png(img_data["bytes"])

                                    # NO junk filter for Chemistry
                                    # Reviewer decides in dashboard

                                    zone = "solution" if is_sol_para else "question"
                                    fn = (f"q{q_num}_{zone}_"
                                          f"{img_data['file'].replace('/', '_')}.png")
                                    gcs_url = upload_to_gcs(img_bytes, fn, paper_id)
                                    url_map[key][zone].append(gcs_url)
                                    images_uploaded += 1
                                    print(f"    Q{q_num} {zone} [DUMP]: "
                                          f"{img_data['file'].split('/')[-1]}")

                else:
                    # ══════════════════════════════════════════
                    # ZONE MODE (Biology) — v4 logic unchanged
                    # Element-order zone detection
                    # Junk filter enabled
                    # ══════════════════════════════════════════

                    current_zone = "question"
                    assignments = []
                    seen_option = False
                    first_option_num = None
                    option_markers = set()
                    has_question_text = False

                    for pi in range(q_start, q_end + 1):
                        if pi >= len(doc.paragraphs):
                            break
                        para = doc.paragraphs[pi]

                        is_sol_para = para_has_sol(para)
                        if is_sol_para:
                            for child in para._element.iter():
                                if child.tag.endswith('}blip'):
                                    rid = child.get(docx_qn('r:embed'))
                                    if (rid and rid not in watermark_rids
                                            and rid in rid_to_image):
                                        assignments.append((rid, "solution"))
                            current_zone = "solution"
                            continue

                        for child in para._element.iter():
                            if child.tag.endswith('}t'):
                                txt = (child.text or '').strip()
                                if not txt:
                                    continue
                                if is_metadata(txt):
                                    continue
                                for mm in re.finditer(r'\(([1-4])\)', txt):
                                    opt = int(mm.group(1))
                                    if not seen_option:
                                        first_option_num = opt
                                    seen_option = True
                                    option_markers.add(opt)
                                    current_zone = f"option_{opt}"
                                if (not seen_option and len(txt) > 10
                                        and not txt.startswith('(')):
                                    has_question_text = True

                            elif child.tag.endswith('}blip'):
                                rid = child.get(docx_qn('r:embed'))
                                if (rid and rid not in watermark_rids
                                        and rid in rid_to_image):
                                    zone = current_zone
                                    if (zone == "question"
                                            and has_question_text
                                            and not seen_option):
                                        zone = "option_1_pending"
                                    assignments.append((rid, zone))

                    if not assignments:
                        continue

                    # Resolve pending option_1
                    final = []
                    for rid, zone in assignments:
                        if zone == "option_1_pending":
                            if first_option_num and first_option_num >= 2:
                                zone = "option_1"
                            else:
                                zone = "question"
                        final.append((rid, zone))
                    assignments = final

                    # Redistribution
                    q_rids = [r for r, z in assignments if z == "question"]
                    filled = {int(z.split('_')[1])
                              for _, z in assignments
                              if z.startswith("option_")}
                    empty = sorted(option_markers - filled)

                    if (len(q_rids) > 0
                            and len(q_rids) == len(empty)
                            and len(empty) >= 3):
                        redist = {}
                        for idx, opt in enumerate(empty):
                            if idx < len(q_rids):
                                redist[q_rids[idx]] = f"option_{opt}"
                        assignments = [(r, redist.get(r, z))
                                       for r, z in assignments]
                        print(f"    Q{q_num}: redistributed "
                              f"{len(redist)} → options {empty}")

                    # Upload with junk filter
                    for rid, zone in assignments:
                        img_data = rid_to_image[rid]
                        img_bytes = ensure_png(img_data["bytes"])

                        if is_junk_image(img_bytes):
                            junk_filtered += 1
                            print(f"    Q{q_num} JUNK: "
                                  f"{img_data['file'].split('/')[-1]}")
                            continue

                        fn = (f"q{q_num}_{zone}_"
                              f"{img_data['file'].replace('/', '_')}.png")
                        gcs_url = upload_to_gcs(img_bytes, fn, paper_id)
                        url_map[key][zone].append(gcs_url)
                        images_uploaded += 1
                        print(f"    Q{q_num} {zone}: "
                              f"{img_data['file'].split('/')[-1]}")

        # ── Summary ──
        print(f"\n{'='*60}")
        print(f"COMPLETE: {images_uploaded} uploaded, "
              f"{junk_filtered} junk filtered")
        for sec in DOCX_SECTIONS:
            keys = [k for k in url_map if k.startswith(sec)]
            if keys:
                nums = sorted(int(k.split('_')[-1]) for k in keys)
                zones = set()
                for k in keys:
                    zones.update(url_map[k].keys())
                mode = "DUMP" if sec in DUMP_SECTIONS else "ZONE"
                print(f"  {sec} ({mode}): {len(keys)} questions "
                      f"(Q{min(nums)}-Q{max(nums)}) "
                      f"zones: {sorted(zones)}")
        print(f"{'='*60}\n")

        return dict(url_map)

    finally:
        try:
            os.unlink(docx_path)
        except Exception:
            pass