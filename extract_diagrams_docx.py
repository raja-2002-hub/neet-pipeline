"""
extract_diagrams_docx.py — DOCX diagram extraction for Chemistry & Biology

v3 — April 2026

STRATEGY:
  The old PDF extractor (extract_diagrams.py) works perfectly for Physics
  because Physics has "Question No." markers. But Chemistry has 0 markers
  (iLovePDF stripped them), causing zone mapping failures.

  This module extracts diagrams ONLY from Chemistry and Biology sections
  in the DOCX file. Physics is left to the PDF extractor.

  main.py merges both url_maps:
    pdf_url_map  = extract_diagrams(pdf_bytes, ...)        # Physics
    docx_url_map = extract_diagrams_docx(docx_bytes, ...)  # Chemistry + Biology
    url_map = {**pdf_url_map, **docx_url_map}

FIXES:
  1. Chemistry question boundary: "Question Type: NEET" / "Difficulty"
  2. Multi-option paragraph image distribution
  3. Junk image pixel filtering (empty rectangles, dots)
  4. Zone tracking after Sol. (no premature reset)
  5. Option vs question image assignment
"""

import re
import io
import zipfile
import tempfile
import os
from collections import defaultdict

from docx import Document
from docx.oxml.ns import qn as docx_qn
from PIL import Image
import numpy as np
from google.cloud import storage

PROJECT         = "project-3639c8e1-b432-4a18-99f"
DIAGRAMS_BUCKET = f"{PROJECT}-diagrams"

MIN_IMAGE_BYTES      = 200
WATERMARK_THRESHOLD  = 10
JUNK_WHITE_THRESHOLD = 245
JUNK_MIN_DARK_PIXELS = 100
JUNK_MIN_DARK_RATIO  = 0.005

# Sections to process from DOCX (Physics handled by PDF extractor)
DOCX_SECTIONS = {"Chemistry", "Biology"}

# Metadata lines — not question content
SKIP_PREFIXES = [
    'Difficulty', 'Topic', 'Concept', 'Single Correct',
    'Expected', 'NEET UG', 'Biology 2016', 'Physics',
    'Question No.', 'Question Type',
]


# ─────────────────────────────────────────────────────────
# GCS Upload
# ─────────────────────────────────────────────────────────

def upload_to_gcs(image_bytes, filename, paper_id):
    sc = storage.Client(project=PROJECT)
    bp = f"{paper_id}/{filename}"
    sc.bucket(DIAGRAMS_BUCKET).blob(bp).upload_from_string(
        image_bytes, content_type="image/png")
    return f"gs://{DIAGRAMS_BUCKET}/{bp}"


# ─────────────────────────────────────────────────────────
# Image Processing
# ─────────────────────────────────────────────────────────

def ensure_png(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode == 'RGBA':
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        out = io.BytesIO()
        img.save(out, format='PNG', optimize=True)
        return out.getvalue()
    except Exception:
        return image_bytes


def is_junk_image(image_bytes):
    """
    FIX 3: Check actual pixel content, not just byte size.
    Filters empty rectangles (Q147/Q170/Q173 solution artifacts),
    tiny dots (Q168), and near-blank images.
    """
    try:
        img = Image.open(io.BytesIO(image_bytes)).convert('RGB')
        arr = np.array(img)
        total_pixels = arr.shape[0] * arr.shape[1]

        # Tiny images are artifacts
        if arr.shape[0] < 10 or arr.shape[1] < 10:
            return True

        brightness = arr.mean(axis=2)
        dark_pixels = int(np.sum(brightness < JUNK_WHITE_THRESHOLD))

        if dark_pixels < JUNK_MIN_DARK_PIXELS:
            return True
        if dark_pixels / total_pixels < JUNK_MIN_DARK_RATIO:
            return True

        # Empty rectangle: only border is dark, interior is white
        if total_pixels > 400 and arr.shape[0] > 6 and arr.shape[1] > 6:
            interior = arr[3:-3, 3:-3]
            interior_dark = int(np.sum(interior.mean(axis=2) < JUNK_WHITE_THRESHOLD))
            interior_total = interior.shape[0] * interior.shape[1]
            if interior_total > 0 and interior_dark / interior_total < 0.001:
                return True

        return False
    except Exception:
        return False


# ─────────────────────────────────────────────────────────
# DOCX Analysis
# ─────────────────────────────────────────────────────────

def detect_watermark_rids(doc):
    rid_count = defaultdict(int)
    for para in doc.paragraphs:
        seen = set()
        for blip in para._element.findall('.//' + docx_qn('a:blip')):
            rid = blip.get(docx_qn('r:embed'))
            if rid and rid not in seen:
                rid_count[rid] += 1
                seen.add(rid)
    wm = {rid for rid, c in rid_count.items() if c >= WATERMARK_THRESHOLD}
    if wm:
        print(f"  Watermark rIds (appear {WATERMARK_THRESHOLD}+ times): {wm}")
    return wm


def build_rid_to_image(docx_path):
    rid_to_image = {}
    with zipfile.ZipFile(docx_path) as z:
        rels = z.read("word/_rels/document.xml.rels").decode()
        rid_to_file = {}
        for m in re.finditer(r'Id="(rId\d+)".*?Target="(media/[^"]+)"', rels):
            rid_to_file[m.group(1)] = m.group(2)
        for rid, media_path in rid_to_file.items():
            try:
                data = z.read(f"word/{media_path}")
                if len(data) >= MIN_IMAGE_BYTES:
                    rid_to_image[rid] = {"bytes": data, "file": media_path}
            except Exception:
                pass
    print(f"  Valid images: {len(rid_to_image)} (filtered spacers < {MIN_IMAGE_BYTES}b)")
    return rid_to_image


def get_para_images(para, watermark_rids):
    rids = []
    for blip in para._element.findall('.//' + docx_qn('a:blip')):
        rid = blip.get(docx_qn('r:embed'))
        if rid and rid not in watermark_rids:
            rids.append(rid)
    return rids


# ─────────────────────────────────────────────────────────
# Section Detection
# ─────────────────────────────────────────────────────────

def detect_section_change(text, current_section):
    """Strict section detection — prevents false matches like 'Topic: Inorganic Chemistry'."""
    if text == "Physics":
        return "Physics"
    if "NEET" in text and "Chemistry" in text and len(text) < 80:
        return "Chemistry"
    if re.match(r'^Biology\s*\d*$', text):
        return "Biology"
    return current_section


# ─────────────────────────────────────────────────────────
# Question Boundary Detection
# ─────────────────────────────────────────────────────────

def is_question_boundary(text):
    """
    FIX 1: Detects question boundaries even without "Question No.".
    Chemistry has 0 "Question No." markers (iLovePDF stripped them).
    Uses "Question Type:" and "Difficulty of question" as alternatives.
    """
    if re.match(r'^Question No\.?\s*\d+', text, re.IGNORECASE):
        return True
    if text.startswith('Question Type:'):
        return True
    if text.startswith('Difficulty of question'):
        return True
    return False


def is_metadata_line(text):
    return any(text.startswith(p) for p in SKIP_PREFIXES)


# ─────────────────────────────────────────────────────────
# Zone Detection
# ─────────────────────────────────────────────────────────

def parse_zone(text):
    """
    FIX 2: Handles multi-option paragraphs like "(1) text (2) (3) (4)".
    Returns (zone, [other_option_numbers]) or (None, None).
    """
    text = text.strip()
    if re.match(r'^Sol\.', text):
        return "solution", []
    option_matches = list(re.finditer(r'\(([1-4])\)', text))
    if option_matches:
        first_opt = int(option_matches[0].group(1))
        later_opts = [int(m.group(1)) for m in option_matches[1:]]
        return f"option_{first_opt}", later_opts
    return None, None


# ─────────────────────────────────────────────────────────
# Main Extraction — Chemistry & Biology only
# ─────────────────────────────────────────────────────────

def extract_diagrams_docx(docx_bytes, paper_id, corrected_questions=None):
    """
    Extract diagrams from DOCX for Chemistry and Biology sections ONLY.
    Physics is skipped — the PDF extractor handles it better because
    Physics has "Question No." markers that give reliable zone mapping.

    Returns url_map: {"Chemistry_144": {"option_2": ["gs://..."]}}
    """
    print(f"\n{'='*60}")
    print(f"DOCX EXTRACTION v3 (Chemistry+Biology) — {paper_id}")
    print(f"{'='*60}")

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.write(docx_bytes)
    tmp.close()
    docx_path = tmp.name

    try:
        doc = Document(docx_path)
        print(f"  Paragraphs: {len(doc.paragraphs)}")

        # ── Step 1: Detect watermarks ──
        watermark_rids = detect_watermark_rids(doc)

        # ── Step 2: Build image map ──
        rid_to_image = build_rid_to_image(docx_path)

        # ── Step 3: Build per-section question lists ──
        section_questions = {"Physics": [], "Chemistry": [], "Biology": []}
        if corrected_questions:
            for q in corrected_questions:
                sec = q.get("section", "")
                if sec in section_questions:
                    section_questions[sec].append(q)
            for sec in section_questions:
                section_questions[sec].sort(
                    key=lambda q: q.get("question_number", 0))
            for sec, qs in section_questions.items():
                if qs:
                    nums = [q["question_number"] for q in qs]
                    tag = "→ DOCX" if sec in DOCX_SECTIONS else "→ SKIP (PDF handles)"
                    print(f"  {sec}: {len(qs)} questions "
                          f"(Q{min(nums)}-Q{max(nums)}) {tag}")

        # ── Step 4: Walk paragraphs ──
        print("\n  Walking paragraphs...")

        url_map = defaultdict(lambda: defaultdict(list))
        current_section = "Physics"
        current_zone = "unknown"

        section_sol = {"Physics": 0, "Chemistry": 0, "Biology": 0}
        section_q_idx = {"Physics": -1, "Chemistry": -1, "Biology": -1}

        # FIX 4: expecting_question — zone stays "solution" after Sol.
        # until a real question boundary is detected
        expecting_question = True

        images_found = 0
        junk_filtered = 0
        physics_skipped = 0

        for pi, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # ── Section detection ──
            new_section = detect_section_change(text, current_section)
            if new_section != current_section:
                print(f"    P{pi}: Section → {new_section}")
                current_section = new_section
                expecting_question = True
                current_zone = "unknown"

            # ── SKIP Physics — let PDF extractor handle it ──
            if current_section not in DOCX_SECTIONS:
                # Still track Sol. count for Physics so the counter
                # stays synchronized, but don't extract images
                if re.match(r'^Sol\.', text):
                    section_sol[current_section] += 1
                    expecting_question = True
                elif is_question_boundary(text):
                    section_q_idx[current_section] = \
                        section_sol[current_section]
                    expecting_question = False

                rids = get_para_images(para, watermark_rids)
                rids = [r for r in rids if r in rid_to_image]
                if rids:
                    physics_skipped += len(rids)
                continue

            # ── Zone detection (Chemistry & Biology) ──

            # Solution marker
            if re.match(r'^Sol\.', text):
                current_zone = "solution"
                section_sol[current_section] += 1
                expecting_question = True
                continue

            # FIX 1: Question boundary
            if is_question_boundary(text):
                section_q_idx[current_section] = \
                    section_sol[current_section]
                expecting_question = False
                current_zone = "question"
                continue

            # Option markers
            zone_result, multi_opts = parse_zone(text)
            if zone_result and zone_result.startswith("option_"):
                # FIX 5: If expecting question and hit (1), question
                # text already passed — advance question index
                if expecting_question:
                    section_q_idx[current_section] = \
                        section_sol[current_section]
                    expecting_question = False
                current_zone = zone_result
            elif zone_result:
                current_zone = zone_result
            else:
                # Content paragraph — if expecting question and this
                # is real text (not metadata), it's question content
                if expecting_question and text and len(text) > 15:
                    if not is_metadata_line(text):
                        section_q_idx[current_section] = \
                            section_sol[current_section]
                        current_zone = "question"
                        expecting_question = False

            # ── Image extraction ──
            rids = get_para_images(para, watermark_rids)
            rids = [r for r in rids if r in rid_to_image]

            if not rids or current_zone == "unknown":
                continue

            # Map to question number
            q_idx = section_q_idx[current_section]
            sec_qs = section_questions.get(current_section, [])
            if q_idx < 0 or q_idx >= len(sec_qs):
                continue

            q_num = sec_qs[q_idx]["question_number"]
            q_sec = sec_qs[q_idx]["section"]
            key = f"{q_sec}_{q_num}"

            # ── FIX 2: Multi-option paragraph distribution ──
            if multi_opts and len(rids) > 0:
                opt_matches = list(re.finditer(r'\(([1-4])\)', text))
                all_opt_nums = [int(m.group(1)) for m in opt_matches]

                # Check if first option has text content
                if len(opt_matches) >= 2:
                    text_between = text[opt_matches[0].end():
                                       opt_matches[1].start()].strip()
                    if text_between:
                        img_zones = [f"option_{n}" for n in all_opt_nums[1:]]
                    else:
                        img_zones = [f"option_{n}" for n in all_opt_nums]
                else:
                    img_zones = [current_zone]

                for idx, rid in enumerate(rids):
                    img_data = rid_to_image[rid]
                    img_bytes = ensure_png(img_data["bytes"])

                    if is_junk_image(img_bytes):
                        junk_filtered += 1
                        print(f"    Q{q_num} JUNK: {img_data['file']}")
                        continue

                    zone = (img_zones[idx]
                            if idx < len(img_zones)
                            else current_zone)

                    fn = (f"q{q_num}_{zone}_"
                          f"{img_data['file'].replace('/', '_')}.png")
                    gcs_url = upload_to_gcs(img_bytes, fn, paper_id)
                    url_map[key][zone].append(gcs_url)
                    images_found += 1
                    print(f"    Q{q_num} {zone}: {img_data['file']}")
            else:
                # Standard: all images → current zone
                for rid in rids:
                    img_data = rid_to_image[rid]
                    img_bytes = ensure_png(img_data["bytes"])

                    if is_junk_image(img_bytes):
                        junk_filtered += 1
                        print(f"    Q{q_num} JUNK: {img_data['file']}")
                        continue

                    fn = (f"q{q_num}_{current_zone}_"
                          f"{img_data['file'].replace('/', '_')}.png")
                    gcs_url = upload_to_gcs(img_bytes, fn, paper_id)
                    url_map[key][current_zone].append(gcs_url)
                    images_found += 1
                    print(f"    Q{q_num} {current_zone}: {img_data['file']}")

        # ── Summary ──
        print(f"\n{'='*60}")
        print(f"DOCX EXTRACTION COMPLETE:")
        print(f"  Chemistry+Biology images: {images_found}")
        print(f"  Physics images skipped:   {physics_skipped} (PDF handles)")
        print(f"  Junk filtered:            {junk_filtered}")
        print(f"  Sol. per section:         {dict(section_sol)}")
        print(f"  Questions with images:    {len(url_map)}")

        for sec in ["Chemistry", "Biology"]:
            sec_keys = [k for k in url_map if k.startswith(sec)]
            if sec_keys:
                q_nums = sorted([int(k.split('_')[-1]) for k in sec_keys])
                zones = set()
                for k in sec_keys:
                    zones.update(url_map[k].keys())
                print(f"  {sec}: {len(sec_keys)} questions "
                      f"(Q{min(q_nums)}-Q{max(q_nums)}) "
                      f"zones: {sorted(zones)}")
        print(f"{'='*60}\n")

        return dict(url_map)

    finally:
        try:
            os.unlink(docx_path)
        except Exception:
            pass